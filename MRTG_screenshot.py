import os.path
import time
import requests
import io
import base64
import shutil
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import timedelta, datetime
from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities


def zabbix_screenshot():
    try:
        if not os.path.isdir('zabbix_screenshot'):
            os.mkdir('zabbix_screenshot')

        new_path = 'zabbix_screenshot/'

        zabbix_url = 'http://125.227.81.68:8080'

        driver.get(zabbix_url)
        time.sleep(1)
        print('zabbix登入中...')
        driver.find_element_by_id('name').send_keys('MIS')
        driver.find_element_by_id('password').send_keys('MisMis28343196')
        driver.find_element_by_id('enter').click()
        time.sleep(1)
        print('登入成功!')
        driver.get('http://125.227.81.68:8080/screens.php?elementid=39')
        time.sleep(1)

        driver.save_screenshot(new_path + 'zabbix 122.147.104.4.png')
        box = (77, 421, 578, 927)  # left, top, right, bottom
        im = Image.open(new_path + 'zabbix 122.147.104.4.png')
        im = im.crop(box)
        im.save(new_path + 'zabbix 122.147.104.4.png')

        print('zabbix 122.147.104.4 cpu & ram screenshot success')
        print('zabbix截圖完成')
    except:
        print('zabbix截圖失敗')


def seednet_get_screenshot():
    try:
        if not os.path.isdir('seednet_screenshot'):
            os.mkdir('seednet_screenshot')

        new_path = 'seednet_screenshot/'

        numbers = len(seednet_LoginId)
        CompId = '24438615'

        for i in range(numbers):
            if i == numbers - 1:
                CompId = '28343196'

            driver.get('https://eservice.seed.net.tw/servlet/project.authsp.LoginServlet?CompId=' + CompId + '&LoginId='
                       + seednet_LoginId[i] + '&Passwd=' + seednet_Password[i])
            driver.get('https://eservice.seed.net.tw/servlet/project.authsp.CookieServlet?URLKey=' + seednet_key[i])
            time.sleep(1)
            driver.save_screenshot(new_path + seednet_LoginId[i] + '.png')

            box = (5, 175, 505, 385)  # left, top, right, bottom

            im = Image.open(new_path + seednet_LoginId[i] + '.png')
            im = im.crop(box)
            im.save(new_path + seednet_LoginId[i] + '.png')

            print(seednet_LoginNanme[i] + ' screenshot success')
        print('遠傳截圖完成')
    except:
        print('遠傳截圖失敗')


def aptg_get_screenshot():
    try:
        if not os.path.isdir('aptg_screenshot'):
            os.mkdir('aptg_screenshot')

        new_path = 'aptg_screenshot/'

        driver.get('https://biz2018.aptg.com.tw/mrtg-frontend/')
        print('登入亞太IDC中!')

        time.sleep(1)
        driver.find_element_by_name('accountType').click()
        driver.find_element_by_id('username').send_keys('CUS-04120800129')
        driver.find_element_by_id('password').send_keys('cus0412')

        driver.save_screenshot(new_path + 'aptg_Verification_code.png')
        box = (1439, 378, 1563, 421)  # left, top, right, bottom
        im = Image.open(new_path + 'aptg_Verification_code.png')
        im = im.crop(box)
        im.save(new_path + 'aptg_Verification_code.png')

        print('驗證碼獲取成功!')
        im_path = new_path + 'aptg_Verification_code.png'

        print('驗證碼辨識中...')
        an = parse_captcha(im_path)

        try:
            print('驗證碼為:' + an)
            driver.find_element_by_id('inCaptcha').send_keys(an)
            driver.find_element_by_id('loginButton').click()
            time.sleep(1)
            try:
                time.sleep(1)
                driver.find_element_by_link_text('流量管理').click()
                time.sleep(1)
                print('登入成功!')
                driver.find_element_by_link_text('數據服務').click()
                time.sleep(2)
                driver.find_element_by_link_text('網際網路').click()
                time.sleep(1)
                driver.find_element_by_css_selector('#\\39 610 .glyphicon-search').click()
                time.sleep(2)
                driver.save_screenshot(new_path + 'VLCK-13121700001.png')
                im = Image.open(new_path + 'VLCK-13121700001.png')
                im = im.crop((282, 690, 783, 895))  # left, top, right, bottom
                im.save(new_path + 'VLCK-13121700001.png')
                print('亞太 VLCK-13121700001 124.219.9網段 screenshot success')

                select = Select(driver.find_element_by_name('contractType'))
                select.select_by_value("9609")
                time.sleep(1)
                driver.find_element_by_css_selector('.btn-primary').click()
                time.sleep(2)
                driver.save_screenshot(new_path + 'VLCK-11122100001.png')
                im = Image.open(new_path + 'VLCK-11122100001.png')
                im = im.crop((282, 690, 783, 895))  # left, top, right, bottom
                im.save(new_path + 'VLCK-11122100001.png')
                print('亞太 VLCK-11122100001 210.202.89.208-223(20Mbits/s) screenshot success')

                select.select_by_value("9608")
                time.sleep(1)
                driver.find_element_by_css_selector('.btn-primary').click()
                time.sleep(2)
                driver.save_screenshot(new_path + 'VLCK-11090800001.png')
                im = Image.open(new_path + 'VLCK-11090800001.png')
                im = im.crop((282, 690, 783, 895))  # left, top, right, bottom
                im.save(new_path + 'VLCK-11090800001.png')
                print('亞太 VLCK-11090800001 60.244.103.80-95(20Mbits/s) screenshot success')

                driver.find_element_by_link_text('系統登出').click()
            except:
                print('亞太登入失敗')
        except:
            print('亞太驗證失敗')
    except:
        print('亞太截圖失敗')

def colo_get_screenshot():
    try:
        if not os.path.isdir('colo_screenshot'):
            os.mkdir('colo_screenshot')

        new_path = 'colo_screenshot/'

        colo_count = len(colo_LoginId)
        colo_url = 'http://61.63.137.92/cacti/index.php'

        driver.get(colo_url)
        time.sleep(1)
        print('宏遠IDC登入中...')
        driver.find_element_by_id('login_username').send_keys('9963')
        driver.find_element_by_id('login_password').send_keys('24438615')
        driver.find_element_by_css_selector('.ui-button').click()
        time.sleep(1)
        print('登入成功!')
        driver.find_element_by_id('tree_anchor-33_anchor').click()
        time.sleep(1)

        now = int(time.time())
        for i in range(colo_count):
            driver.get('http://61.63.137.92/cacti/graph_image.php?local_graph_id=' + colo_graph_id[i]
                       + '&graph_start=' + str(int(now-86400)) + '&graph_end=' + str(now)
                       + '&graph_width=500&graph_height=120')
            time.sleep(1)
            driver.save_screenshot(new_path + colo_LoginId[i] + '.png')
            box = (663, 409, 1258, 672)  # left, top, right, bottom
            im = Image.open(new_path + colo_LoginId[i] + '.png')
            im = im.crop(box)
            im.save(new_path + colo_LoginId[i] + '.png')

            print(colo_LoginId[i] + ' screenshot success')
        print('宏遠截圖完成')
    except:
        print('宏遠截圖失敗')


def sonet_screenshot():
    try:
        if not os.path.isdir('sonet_screenshot'):
            os.mkdir('sonet_screenshot')

        new_path = 'sonet_screenshot/'

        sonet_count = len(sonet_LoginId)

        for i in range(sonet_count):
            driver.get(sonet_url[i])
            time.sleep(1)
            driver.save_screenshot(new_path + sonet_LoginId[i] + '.png')

            box = (5, 318, 543, 502)  # left, top, right, bottom

            im = Image.open(new_path + sonet_LoginId[i] + '.png')
            im = im.crop(box)
            im.save(new_path + sonet_LoginId[i] + '.png')

            print(sonet_LoginId[i] + ' screenshot success')
        print('台北SONET截圖完成')
    except:
        print('台北SONET截圖失敗')


def hk_screenshot():
    try:
        if not os.path.isdir('hk_screenshot'):
            os.mkdir('hk_screenshot')

        new_path = 'hk_screenshot/'
        print('登入hk IDC中!')
        driver.get('https://mc.citictel-cpc.com/web/guest')
        time.sleep(1)
        driver.find_element_by_name('_58_gid').send_keys('TW43497')
        driver.find_element_by_name('_58_login').send_keys('TW43497')
        driver.find_element_by_id('gidPw').send_keys(r'Wan@24522811')
        driver.find_element_by_link_text('Login').click()
        print('登入成功!')
        print('擷取流量圖中!')

        def IsElementExist(element_link_text):
            try:
                driver.find_element_by_link_text(element_link_text)
                return True
            except:
                return False

        WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.LINK_TEXT, 'Skip'))
        )

        if IsElementExist('Skip'):
            driver.find_element_by_link_text('Skip').click()

        element = WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.ID, 'first_item_2'))
        )
        element.click()
        driver.switch_to.frame(1)

        try:
            element = WebDriverWait(driver, 30).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, 'tr:nth-child(1) a:nth-child(2) > img'))
            )
            element.click()
        except:
            print('new2載入錯誤，重新執行')
            element = WebDriverWait(driver, 30).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, 'tr:nth-child(1) a:nth-child(2) > img'))
            )
            element.click()

        time.sleep(8)
        driver.switch_to.parent_frame()
        time.sleep(1)
        js_400 = 'var q=document.documentElement.scrollTop=400'
        driver.execute_script(js_400)
        time.sleep(2)

        driver.save_screenshot(new_path + 'TW43497-new2.png')
        box = (285, 80, 1650, 880)  # left, top, right, bottom
        im = Image.open(new_path + 'TW43497-new2.png')
        im = im.crop(box)
        im.save(new_path + 'TW43497-new2.png')
        print('TW43497-new2 screenshot success')

        time.sleep(1)
        element = WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.ID, 'first_item_2'))
        )
        element.click()
        driver.switch_to.frame(1)

        try:
            element = WebDriverWait(driver, 30).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, 'tr:nth-child(2) a:nth-child(2) > img'))
            )
            element.click()
        except:
            print('new1載入錯誤，重新執行')
            element = WebDriverWait(driver, 30).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, 'tr:nth-child(2) a:nth-child(2) > img'))
            )
            element.click()

        time.sleep(8)
        driver.switch_to.parent_frame()
        time.sleep(1)
        js_400 = 'var q=document.documentElement.scrollTop=400'
        driver.execute_script(js_400)
        time.sleep(2)

        driver.save_screenshot(new_path + 'TW43497-new1.png')
        box = (285, 80, 1650, 880)  # left, top, right, bottom
        im = Image.open(new_path + 'TW43497-new1.png')
        im = im.crop(box)
        im.save(new_path + 'TW43497-new1.png')
        print('TW43497-new1 screenshot success')
        print('香港IDC截圖完成')
    except:
        print('香港IDC截圖失敗')


def waninbank_screenshot():
    try:
        waninbank = 'http://mrtg.waninbank.tw'
        if not os.path.isdir('waninbank_screenshot'):
            os.mkdir('waninbank_screenshot')

        new_path = 'waninbank_screenshot/'
        print('登入中軟IDC中!')
        driver.get('http://mrtg.waninbank.tw/index.php')
        time.sleep(1)
        driver.find_element_by_name('login_username').send_keys('WS0000001')
        driver.find_element_by_name('login_password').send_keys('24438615')
        driver.find_element_by_xpath("//input[@value='登录']").click()
        print('登入成功!')
        now = int(time.time())

        for i in range(len(waninbank_LoginId)):
            driver.get(waninbank + '/graph_image.php?local_graph_id=' + waninbank_urlcode[i]
                       + '&rra_id=0&view_type=tree&' + 'graph_start=' + str(int(now-86400)) + '&graph_end=' + str(now))
            time.sleep(1)
            driver.save_screenshot(new_path + waninbank_LoginId[i] + '.png')

            box = (667, 421, 1254, 659)  # left, top, right, bottom

            im = Image.open(new_path + waninbank_LoginId[i] + '.png')
            im = im.crop(box)
            im.save(new_path + waninbank_LoginId[i] + '.png')

            print(waninbank_LoginId[i] + ' screenshot success')
        print('中軟IDC截圖完成')
    except:
        print('中軟IDC截圖失敗')




def digicentre_screenshot():
    try:
        if not os.path.isdir('digicentre_screenshot'):
            os.mkdir('digicentre_screenshot')

        new_path = 'digicentre_screenshot/'

        print('登入果核IDC中!')
        driver.get('https://tw-sslvpn.digicentre.com/dana-na/auth/url_default/welcome.cgi')
        time.sleep(1)
        driver.find_element_by_id('username').send_keys('wanin')
        driver.find_element_by_id('password').send_keys('Wanin2016')
        driver.find_element_by_id('btnSubmit_6').click()

        time.sleep(1)

        def IsElementExist(element_id):
            try:
                driver.find_element_by_id(element_id)
                return True
            except:
                return False

        if IsElementExist('btnContinue'):
            driver.find_element_by_id('btnContinue').click()

        element = driver.find_element_by_css_selector('#table_webbookmarkline_2 b')
        ActionChains(driver).key_down(Keys.CONTROL).click(element).key_up(Keys.CONTROL).perform()

        nowhandle = driver.current_window_handle
        allhandles = driver.window_handles
        for handle in allhandles:
            if handle != nowhandle:
                driver.switch_to.window(handle)
                time.sleep(1)

        driver.find_element_by_name('login_username').send_keys('wanin')
        driver.find_element_by_name('login_password').send_keys('Wanin2016')
        driver.find_element_by_css_selector('td:nth-child(1) > input').click()
        print('登入成功!')

        time.sleep(1)
        driver.save_screenshot(new_path + '網銀國際A.png')
        box = (740, 205, 1345, 480)  # left, top, right, bottom
        im = Image.open(new_path + '網銀國際A.png')
        im = im.crop(box)
        im.save(new_path + '網銀國際A.png')
        print('網銀國際A screenshot success')

        driver.save_screenshot(new_path + '網銀國際B.png')
        box = (740, 485, 1345, 758)  # left, top, right, bottom
        im = Image.open(new_path + '網銀國際B.png')
        im = im.crop(box)
        im.save(new_path + '網銀國際B.png')
        print('網銀國際B screenshot success')

        js_bottom = "var q=document.documentElement.scrollTop=10000"
        driver.execute_script(js_bottom)
        time.sleep(1)

        driver.save_screenshot(new_path + '網銀國際D.png')
        box = (740, 687, 1345, 1010)  # left, top, right, bottom
        im = Image.open(new_path + '網銀國際D.png')
        im = im.crop(box)
        im.save(new_path + '網銀國際D.png')
        print('網銀國際D screenshot success')

        driver.find_element_by_link_text('Logout').click()
        driver.close()
        driver.switch_to.window(nowhandle)
        driver.find_element_by_id('imgNavSignOut').click()
        print('果核IDC截圖完成')
    except:
        print('果核IDC截圖失敗')



def hinet_get_screenshot():
    try:
        if not os.path.isdir('hinet_screenshot'):
            os.mkdir('hinet_screenshot')

        new_path = 'hinet_screenshot/'

        driver.get('https://enoc.hinet.net/eNoc/login.do')
        print('登入中華電信IDC中!')

        time.sleep(1)
        driver.find_element_by_id('hn_no').send_keys('73898925')
        driver.find_element_by_id('password').send_keys('Wan24522811')

        driver.save_screenshot(new_path + 'hinet_Verification_code.png')
        # box = (241, 334, 333, 356)  # left, top, right, bottom (minisize
        box = (701, 334, 794, 356)  # left, top, right, bottom
        im = Image.open(new_path + 'hinet_Verification_code.png')
        im = im.crop(box)
        im.save(new_path + 'hinet_Verification_code.png')

        print('驗證碼獲取成功!')
        im_path = new_path + 'hinet_Verification_code.png'

        print('驗證碼辨識中...')
        an = parse_captcha(im_path)

        try:
            print('驗證碼為:' + an)
            driver.find_element_by_id('inputCaptcha').send_keys(an)
            driver.find_element_by_css_selector('td > a > img').click()
            time.sleep(1)
            try:
                print('登入成功!')
                print('擷取流量圖中!')
                driver.find_element_by_link_text('電路及時狀態資訊').click()
                driver.find_element_by_css_selector('.m02 > a').click()
                time.sleep(1)
                driver.find_element_by_xpath("//input[@value=\'查詢\']").click()
                time.sleep(1)
                driver.find_element_by_link_text('即時流量').click()
                time.sleep(1)
                driver.find_element_by_link_text('點我').click()
                time.sleep(1)

                nowhandle = driver.current_window_handle
                allhandles = driver.window_handles
                for handle in allhandles:
                    if handle != nowhandle:
                        driver.switch_to.window(handle)
                        time.sleep(1)

                driver.save_screenshot(new_path + '73898925.png')
                im = Image.open(new_path + '73898925.png')
                im = im.crop((5, 232, 571, 444))  # left, top, right, bottom
                im.save(new_path + '73898925.png')
                print('網銀國際 編號：73898925 60.249.42.224-239 (4M/4M) screenshot success')
                driver.close()
                driver.switch_to.window(nowhandle)

                driver.find_element_by_css_selector('#tab2 .back_list > a').click()
                time.sleep(1)
                driver.find_element_by_xpath("(//input[@value=\'查詢\'])[2]").click()
                time.sleep(1)
                driver.find_element_by_link_text('即時流量').click()
                time.sleep(1)
                driver.find_element_by_link_text('點我').click()
                time.sleep(1)

                nowhandle = driver.current_window_handle
                allhandles = driver.window_handles
                for handle in allhandles:
                    if handle != nowhandle:
                        driver.switch_to.window(handle)
                        time.sleep(1)

                driver.save_screenshot(new_path + '73806990.png')
                im = Image.open(new_path + '73806990.png')
                im = im.crop((5, 230, 580, 445))  # left, top, right, bottom
                im.save(new_path + '73806990.png')
                print('宏鑫多媒體 編號：73806990 60.249.40.112-127 (4M/4M) screenshot success')
                driver.close()
                driver.switch_to.window(nowhandle)

                print('中華黎明IDC截圖完成')
            except:
                print('截圖失敗!')
        except:
            print('驗證失敗!')
    except:
        print('中華黎明IDC截圖失敗')


def parse_captcha(image_path):
    # Load image into memory
    buffer = io.BytesIO()
    image = Image.open(image_path)
    image.save(buffer, format='png')
    # Use base64 to encode image buffer
    img_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
    # Anti-captcha API structure
    data = {
        'clientKey': 'ed0953a08c0a0d7b77003015973c95d7',
        'task': {
            'type': 'ImageToTextTask',
            'body': img_str,
            'phrase': False,
            'case': False,
            'numeric': 1,
            'math': 0,
            'minLength': 4,
            'maxLength': 6
        }
    }
    # Create a ImageToTextTask and retrieve taskId from response
    r = requests.post('https://api.anti-captcha.com/createTask', json=data)
    r.raise_for_status()
    task_id = r.json()['taskId']
    # Polling for task finish.
    ret = ""
    while True:
        data = {
            'clientKey': 'ed0953a08c0a0d7b77003015973c95d7',
            'taskId': task_id
        }
        r = requests.post('https://api.anti-captcha.com/getTaskResult', json=data)
        r.raise_for_status()
        if r.json()['status'] == 'ready':
            ret = r.json()['solution']['text']
            break
        time.sleep(1)
    return ret


def big_title(word):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(word)
    run.font.size = Pt(16)
    run.font.name = u'微軟正黑體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
    run.font.bold = True
    run.font.underline = True
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)


def small_title(word):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(word)
    run.font.size = Pt(12)
    run.font.name = u'Calibri'
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)


def add_picture(folder, name, height):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('')
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    try:
        run.add_picture(folder + '_screenshot/' + name + '.png', height=Cm(height))
    except:
        print(name + '.png 截圖失敗，請重新執行或手動截圖!')


def red_word(word):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(word)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)

seednet_LoginId = [
        'CA501056', 'CA501057', 'C9907017', 'C9907018', 'C9908004', 'C9908037', 'CA501058', 'CA205019', 'CA011017',
        'CA112012', 'CA112014', 'CA201016', 'CA103040', 'CA103043', 'CA301018', 'CA305005', 'CA305007', 'CA305008',
        'CA305009', 'CA306012', 'CA308009', 'CA302002', 'CA302003', 'CA302014', 'CA210036', 'CA209020', 'CA209029',
        'CA212025', 'CA205020', 'CA501059', 'CA311012', 'CA405034', 'CA405035', 'CA405036', 'CA502037', 'CA502038',
        'CA502039', 'CA502040', 'CA506096', 'C9911012', 'CA212026', 'CA008001'
    ]


seednet_LoginNanme = [
    '遠傳太平★220.229.226.240-254★C30523-CA501056', '遠傳太平★218.32.166.0-15★C30611-CA501057',
    '遠傳太平★122.147.105.0~15★C9907017', '遠傳太平★122.147.104.0~15★C9907018',
    '遠傳太平★122.147.106.0-15★C9908004', '遠傳太平★122.147.107.0-15★C9908037',
    '遠傳太平★218.32.166.16-23★CA501058', '遠傳太平★175.41.52.193-197★CA205019',
    '遠傳中港★210.64.214.112-119★CA011017', '遠傳中港★210.66.176.48-63★CA112012',
    '遠傳中港★210.244.17.24-31★CA112014', '遠傳中港★210.64.48.0-63★CA201016',
    '遠傳中港★218.32.166.48-55★CA103040', '遠傳中港★220.229.226.80-87★CA103043',
    '遠傳中港★210.64.205.40-47★210.244.17.48-55★CA301018', '遠傳中港★210.66.177.16-23★CA305005',
    '遠傳中港★210.64.205.192-199★CA305007', '遠傳中港★210.64.215.128-135★CA305008',
    '遠傳中港★210.244.17.56-63★CA305009', '遠傳中港★210.64.205.200-207★CA306012',
    '遠傳中港★210.64.214.64-71★CA308009', '遠傳中港★210.64.215.64-71★CA302002',
    '遠傳中港★210.64.216.64-71★CA302003', '遠傳中港★210.64.216.88-95★CA302014',
    '遠傳中港★210.64.214.120-127★210.66.176.128-191★CA210036',
    '遠傳中港★210.64.216.16-31★210.244.25.24-31★CA209020',
    '遠傳中港★210.64.205.120-127★210.64.214.16-31★CA209029', '遠傳中港★210.64.203.0-15★CA212025',
    '遠傳中港★210.244.17.16-23★CA205020', '遠傳中港★210.64.214.96-111★CA501059',
    '遠傳中港★210.66.176.0~7★CA311012', '遠傳中港★61.56.214.132~158★CA405034',
    '遠傳中港★210.244.27.0~31★CA405035', '遠傳中港★210.244.26.64~95★CA405036',
    '遠傳中港★210.244.27.48-63★CA502037', '遠傳中港★210.244.27.64-79★CA502038',
    '遠傳中港★210.244.26.112-127★CA502039', '遠傳中港★210.244.26.128-143★CA502040',
    '遠傳中港★210.64.203.16~31★CA506096', '遠傳台北★113.196.230.0-15★C9911012',
    '遠傳台北★113.196.238.0-15★CA212026', '遠傳雲端★113.196.120.90-93★CA008001'
]

seednet_Password = [
    '78em9c9i', 'f687djyj', 'epckmffk', 'a8tffefu', 'f3wkd7je', '84cv749u', 'j3cyiv4w', 'm879xp9y', 'm9adfd7c',
    'ir3deuwf', 'ii3x4mcr', 'vd4nne3j', '3868rakr', 'vtptu8r6', 'y9a37cfx', 'kfp9d4u4', 'yced7a9x', 'w7f7vec4',
    '4a6ta7tx', '4kvt6paj', 'wimi39rp', 'afd76ju6', '6wxe8rnx', 'vead789p', 'a8e4fmkm', 'epppai8m', '97wt3pta',
    '8cwui6kd', 'rm3mfjk4', 'yxa8ykt4', '7xi4ktj9', 'ixtne98a', '3nnk4mfj', 'iwku4e64', '7fnj9wp8', 'ay39m88d',
    '7k884n9m', 'i3nprkjf', '8fmfc8ac', 'day99ykj', '9mjtj3fa', 'w9nd8vc8'
]

seednet_key = [
    'S200005006--1813304089', 'S200005006--87823192', 'S200005006-1764565594', 'S200005006--804920805',
    'S200005006-565593943', 'S200005006--1291964489', 'S200005006-1637657705', 'S200005006-2003380877',
    'S200005006-1060261394', 'S200005006--177791797', 'S200005006--1021797299', 'S200005006--1114343218',
    'S200005006-482667460', 'S200005006-1364142855', 'S200005006--496768721', 'S200005006--1092295671',
    'S200005006--1936301173', 'S200005006--210820276', 'S200005006-1514660621', 'S200005006-1609333188',
    'S200005006-2118105322', 'S200005006-1717751529', 'S200005006--851734870', 'S200005006--1470921014',
    'S200005006-230769710', 'S200005006--749463649', 'S200005006-1894962536', 'S200005006--179403444',
    'S200005006-1309254947', 'S200005006--931828694', 'S200005006-1112564202', 'S200005006-199736900',
    'S200005006-1925217797', 'S200005006--644268602', 'S200005006-1939347593', 'S200005006--630138806',
    'S200005006-1095342091', 'S200005006-401216161',  'S200005006-1267015038',
    'S200005006--402477284', 'S200005006-1546077453', 'S200005006-1109387333',
    ]

aptg_url = [
        'http://203.79.223.23/mrtg/tcfx_448_yl01/mrtg_tcfx_448_yl01_ethernet8.php',
        'http://203.79.223.23/mrtg/tcfx_448_yl01/mrtg_tcfx_448_yl01_ethernet7.php',
        'http://203.79.223.23/mrtg/tcfi_4802_ld01/mrtg_tcfi_4802_ld01_ethernet6.php'
    ]

colo_graph_id = [
    '585', '92', '6', '528', '586', '532'
]
colo_LoginId = [
        '台中★122.201.145.48-55★9963IAP', '國貿★210.209.15.241★9808', '國貿★210.209.13.192-199★500028',
        '台中★211.73.3.8-15★9963PORT1', '台中★211.73.7.216-223★9963PORT2', '台中★211.73.7.192-198★9963PORT3'
    ]

sonet_url = [
        'http://hongxin:Wan24438615@cus.so-net.net.tw/f13/f13-1.html',
        'http://hongxin:Wan24438615@cus.so-net.net.tw/f13/f13-2.html',
        'http://hongxin:Wan24438615@cus.so-net.net.tw/f13/f13-3.html',
        'http://hongxin:Wan24438615@cus.so-net.net.tw/f13/f13-4.html',
        'http://hongxin:Wan24438615@cus.so-net.net.tw/f13/f13-5.html']

sonet_LoginId = ['台北_宏鑫多媒體1', '台北_宏鑫多媒體2', '台北_宏鑫多媒體3', '台北_宏鑫多媒體4', '台北_宏鑫多媒體5']

waninbank_LoginId = [
    'WS001_103.130.124.128_27_Web Service(1)_100M', 'WS001_103.130.124.128_27_Web Service(2)_100M',
    'WS001_103.130.124.128_27_Web Service(3)_100M', 'WS001_103.130.124.128_27_Web Service(4)_100M',
    'WS002_103.130.127.160_28_Web Service(1)_100M', 'WS002_103.130.127.160_28_Web Service(2)_100M',
    'WS002_103.130.127.160_28_Web Service(3)_100M', 'WS002_103.130.127.160_28_Web Service(4)_100M',
    'WS003_103.130.125.128_28_星城韓文(1)_100M', 'WS003_103.130.125.128_28_星城韓文(2)_100M',
    'WS004_103.130.124.160_27_黑Online(1)_100M', 'WS004_103.130.124.160_27_黑Online(2)_100M',
    'WS005_103.130.124.192_27_測試網段(1)_100M', 'WS005_103.130.124.192_27_測試網段(2)_100M',
    'WS007_103.130.124.224_27_QPP(區塊鏈)(1)_100M', 'WS007_103.130.124.224_27_QPP(區塊鏈)(2)_100M',
    'WS008_103.130.125.0_27_CODE2040(1)_100M', 'WS008_103.130.125.0_27_CODE2040(2)_100M',
    'WS009_103.130.125.144_28_Redis_100M', 'WS010_103.130.127.144_28_Redis_100M',
    'WS011_103.130.127.128_28_星城海外版(日、越)(1)_10M', 'WS011_103.130.127.128_28_星城海外版(日、越)(2)_10M',
    'WS012_103.146.212.128_28_星城海外版(日、越)(1)_10M', 'WS012_103.146.212.128_28_星城海外版(日、越)(2)_10M',
    'WS013_103.130.127.0_25_Fantacity_500M', 'WS013_103.130.127.0_25_Fantacity(2)_500M',
    'WS014_103.146.213.0_24_Fantacity_500M', 'WS014_103.146.213.0_24_Fantacity(2)_500M',
    'WS015_150.116.54.0_25_Fantacity_500M', 'WS015_150.116.54.0_25_Fantacity(2)_500M',
    'WS016_103.130.125.192_26_星城測試環境(1)_10M', 'WS016_103.130.125.192_26_星城測試環境(2)_10M',
    'WS017_103.130.125.48_29_Web Service測試用(1)_10M', 'WS017_103.130.125.48_29_Web Service測試用(2)_10M',
    'WS018_103.130.127.192_27_新molo(1)_100M', 'WS018_103.130.127.192_27_新molo(2)_100M',
    'WS019_103.130.127.224_28_Oinkey Server(1)_100M', 'WS019_103.130.127.224_28_Oinkey Server(2)_100M',
    'WS020_103.130.127.240_28_星城海外版(泰)(1)_60M', 'WS020_103.130.127.240_28_星城海外版(泰)(2)_60M',
    'WS021_103.130.125.32_28_星城鯊(1)_100M', 'WS021_103.130.125.32_28_星城鯊(2)_100M',
    'WS022_103.130.124.40_29_新越南測試區(1)_60M', 'WS022_103.130.124.40_29_新越南測試區(2)_60M',
    'WS023_103.130.124.48_29_舊越南測試區(1)_60M', 'WS023_103.130.124.48_29_舊越南測試區(2)_60M',
    'WS024_103.130.124.56_29_星城海外版(日)(1)_60M', 'WS024_103.130.124.56_29_星城海外版(日)(2)_60M',
    'ADSL_星城下載用_1G_600M'
]

waninbank_urlcode = [
    '843', '895', '847', '899', '844', '896', '848', '900', '1057', '1109', '737', '789', '739', '791',
    '740', '792', '741', '793', '849', '901', '1055', '1107', '1056', '1108', '1058', '1110', '631', '683',
    '1060', '1112', '1061', '1113', '845', '897', '1059', '1111', '1062', '1114', '1063', '1115', '1930',
    '796', '1064', '1116', '1065', '1117', '1066', '1118', '742'
]
if __name__ == '__main__':

    d = DesiredCapabilities.CHROME
    d['loggingPrefs'] = {'performance': 'ALL'}
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--allow-running-insecure-content')
    chrome_options.add_argument('--headless')
    # chrome_options.add_argument('--start-maximized')
    chrome_options.add_argument('--window-size=1920x1080')

    driver = webdriver.Chrome(executable_path='chromedriver.exe', chrome_options=chrome_options, desired_capabilities=d)

    start_time_now = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
    start_time = datetime.now()
    print('開始時間:', start_time_now)

    zabbix_screenshot()
    seednet_get_screenshot()
    aptg_get_screenshot()
    colo_get_screenshot()
    hk_screenshot()
    waninbank_screenshot()
    digicentre_screenshot()
    sonet_screenshot()
    hinet_get_screenshot()

    driver.close()

    # seednet = 遠傳
    # aptg = 亞太
    # colo = 宏遠
    # sonet = 台北sonet
    # waninbank = 中軟IDC
    # hinet = 中華電信 黎明機房
    # digicentre = 果核機房

    doc = Document()
    section = doc.sections[0]
    new_width, new_height = Cm(21), Cm(29.7)
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1)

    big_title('遠傳太平122.147.104.4監控CPU、RAM')

    add_picture('zabbix', 'zabbix 122.147.104.4', 13.5)
    # ---------------------------------遠傳太平IDC機房-------------------------------
    big_title('遠傳太平IDC機房')

    # ---------------------------------CA501056-------------------------------
    small_title('C30523-CA501056 (網銀國際)  220.229.226.240-254 (3M/3M)')
    add_picture('seednet', 'CA501056', 6)

    for i in range(5):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ---------------------------------CA501057-------------------------------
    small_title('C30611-CA501057 (網銀國際)  218.32.166.0-15 (3M/3M)')
    add_picture('seednet', 'CA501057', 6)

    # ---------------------------------C9907017-------------------------------
    small_title('C9907017(網銀國際)  122.147.105.0~15 (3M/3M)')
    add_picture('seednet', 'C9907017', 6)

    # ---------------------------------C9907018-------------------------------
    small_title('C9907018(網銀國際)  122.147.104.0~15 (6M/6M)')
    add_picture('seednet', 'C9907018', 6)

    # ---------------------------------C9908004-------------------------------
    small_title('C9908004(網銀國際)  122.147.106.0-15 (4M/4M)')
    add_picture('seednet', 'C9908004', 6)

    # ---------------------------------C9908037-------------------------------
    small_title('C9908037(網銀國際)  122.147.107.0-15 (3M/3M)')
    add_picture('seednet', 'C9908037', 6)

    # ---------------------------------亞太-太平★60.244.103.80-95★11090800001-------------------------------
    small_title('CUS-04120800129 (網銀國際-)-VLCK-11090800001  60.244.103.80-95(20Mbits/s)')
    red_word('最左邊為最新時間')
    add_picture('aptg', 'VLCK-11090800001', 6)

    # ---------------------------------亞太-太平★210.202.89.208-223★11122100001-------------------------------
    small_title('CUS-04120800129 (網銀國際-)-VLCK-11122100001  210.202.89.208-223(20Mbits/s)')
    red_word('最左邊為最新時間')
    add_picture('aptg', 'VLCK-11122100001', 6)

    for i in range(7):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ---------------------------------亞太-中港★124.219.9.64-79★13121700001-------------------------------
    small_title('CUS-04120800129 (網銀國際-)-VLCK13121700001  124.219.9.網段')
    red_word('最左邊為最新時間')
    add_picture('aptg', 'VLCK-13121700001', 6)

    # ---------------------------------CA501058-------------------------------
    small_title('CA501058 (網銀國際)  218.32.166.16-23')
    add_picture('seednet', 'CA501058', 6)

    # ---------------------------------CA205019-------------------------------
    small_title('CA205019(網銀國際)  175.41.52.193-197')
    add_picture('seednet', 'CA205019', 6)

    for i in range(7):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ---------------------------------遠傳中港IDC機房-------------------------------
    big_title('遠傳中港IDC機房')

    # ---------------------------------CA011017-------------------------------
    small_title('CA011017(網銀國際)  210.64.214.112-119')
    add_picture('seednet', 'CA011017', 6)

    # ---------------------------------CA112012-------------------------------
    small_title('CA112012(網銀國際)  210.66.176.48-63')
    add_picture('seednet', 'CA112012', 6)

    # ---------------------------------CA112014-------------------------------
    small_title('CA112014(網銀國際)  210.244.17.24-31')
    add_picture('seednet', 'CA112014', 6)

    for i in range(7):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ---------------------------------CA201016-------------------------------
    small_title('CA201016(網銀國際)  210.64.48.0-63')
    add_picture('seednet', 'CA201016', 6)

    # ---------------------------------CA103040-------------------------------
    small_title('CA103040(網銀國際)  218.32.166.48-55')
    add_picture('seednet', 'CA103040', 6)

    # ---------------------------------CA103043-------------------------------
    small_title('CA103043(網銀國際)  220.229.226.80-87')
    add_picture('seednet', 'CA103043', 6)

    # ---------------------------------CA301018-------------------------------
    small_title('CA301018(網銀國際)  210.64.205.40-47 ; 210.244.17.48-55')
    add_picture('seednet', 'CA301018', 6)

    # ---------------------------------CA305005-------------------------------
    small_title('CA305005(網銀國際)  210.66.177.16-23')
    add_picture('seednet', 'CA305005', 6)

    # ---------------------------------CA305007-------------------------------
    small_title('CA305007(網銀國際)  210.64.205.192-199')
    add_picture('seednet', 'CA305007', 6)

    # ---------------------------------CA305008-------------------------------
    small_title('CA305008(網銀國際)  210.64.215.128-135')
    add_picture('seednet', 'CA305008', 6)

    # ---------------------------------CA305009-------------------------------
    small_title('CA305009(網銀國際)  210.244.17.56-63')
    add_picture('seednet', 'CA305009', 6)

    # ---------------------------------CA306012-------------------------------
    small_title('CA306012網銀國際)  210.64.205.200-207')
    add_picture('seednet', 'CA306012', 6)

    # ---------------------------------CA308009-------------------------------
    small_title('CA308009 (網銀國際)  210.64.214.64-71')
    add_picture('seednet', 'CA308009', 6)

    # ---------------------------------CA302002-------------------------------
    small_title('CA302002網銀國際)  210.64.215.64-71')
    add_picture('seednet', 'CA302002', 6)

    # ---------------------------------CA302003-------------------------------
    small_title('CA302003(網銀國際)  210.64.216.64-71')
    add_picture('seednet', 'CA302003', 6)

    # ---------------------------------CA302014-------------------------------
    small_title('CA302014(網銀國際)  210.64.216.88-95')
    add_picture('seednet', 'CA302014', 6)

    # ---------------------------------CA210036-------------------------------
    small_title('CA210036(網銀國際)  210.64.214.120-127 ; 210.66.176.128-191')
    add_picture('seednet', 'CA210036', 6)

    # ---------------------------------CA209020-------------------------------
    small_title('CA209020(網銀國際)  210.64.216.16-31 ; 210.244.25.24-31')
    add_picture('seednet', 'CA209020', 6)

    # ---------------------------------CA209029-------------------------------
    small_title('CA209029(網銀國際)  210.64.205.120-127 ; 210.64.214.16-31')
    add_picture('seednet', 'CA209029', 6)

    # ---------------------------------CA212025-------------------------------
    small_title('CA212025(網銀國際)  210.64.203.0-15')
    add_picture('seednet', 'CA212025', 6)

    # ---------------------------------CA205020-------------------------------
    small_title('CA205020(網銀國際)  210.244.17.16-23')
    add_picture('seednet', 'CA205020', 6)

    # ---------------------------------CA501059-------------------------------
    small_title('CA501059 (網銀國際)  210.64.214.96-111')
    add_picture('seednet', 'CA501059', 6)

    # ---------------------------------CA311011-------------------------------
    # small_title('CA311011 (網銀國際)  210.64.205.0~7')
    # add_picture('seednet', 'CA311011', 6)

    # ---------------------------------CA311012-------------------------------
    small_title('CA311012 (網銀國際)  210.66.176.0~7')
    add_picture('seednet', 'CA311012', 6)

    # ---------------------------------CA311013-------------------------------
    # small_title('CA311013 (網銀國際)  210.64.216.96~103')
    # red_word('撤機')

    # ---------------------------------CA405034-------------------------------
    small_title('CA405034 (網銀國際)  61.56.214.132~158')
    add_picture('seednet', 'CA405034', 6)

    # ---------------------------------CA405035-------------------------------
    small_title('CA405035 (網銀國際)  210.244.27.0~31')
    add_picture('seednet', 'CA405035', 6)

    # ---------------------------------CA405036-------------------------------
    small_title('CA405036 (網銀國際)  210.244.26.64~95')
    add_picture('seednet', 'CA405036', 6)

    # ---------------------------------CA502037-------------------------------
    small_title('CA502037 (網銀國際)  210.244.27.48-63')
    add_picture('seednet', 'CA502037', 6)

    # ---------------------------------CA502038-------------------------------
    small_title('CA502038 (網銀國際)  210.244.27.64-79')
    add_picture('seednet', 'CA502038', 6)

    # ---------------------------------CA502039-------------------------------
    small_title('CA502039 (網銀國際)  210.244.26.112-127')
    add_picture('seednet', 'CA502039', 6)

    # ---------------------------------CA502040-------------------------------
    small_title('CA502040 (網銀國際)  210.244.26.128-143')
    add_picture('seednet', 'CA502040', 6)

    # ---------------------------------CA506095-------------------------------
    # small_title('CA506095 (網銀國際)  210.64.48.64~79')
    # add_picture('seednet', 'CA506095', 6) 20200730 撤線

    # ---------------------------------CA506096-------------------------------
    small_title('CA506096 (網銀國際)  210.64.203.16~31')
    add_picture('seednet', 'CA506096', 6)

    # ---------------------------------CA506097-------------------------------
    # small_title('CA506097 (網銀國際)  210.64.205.128~143')
    # red_word('撤機')

    # ---------------------------------CA506098-------------------------------
    # small_title('CA506098 (網銀國際)  210.244.27.33~47')
    # red_word('撤機')

    # ---------------------------------遠傳台北IDC機房-------------------------------
    big_title('遠傳台北IDC機房')

    # ---------------------------------C9911012-------------------------------
    small_title('C9911012(網銀國際)  113.196.230.0-15')
    add_picture('seednet', 'C9911012', 6)

    # ---------------------------------CA212026-------------------------------
    small_title('CA212026(網銀國際)  113.196.238.0-15')
    add_picture('seednet', 'CA212026', 6)

    # ---------------------------------遠傳雲端-------------------------------
    big_title('遠傳雲端')

    # ---------------------------------CA008001-------------------------------
    small_title('CA008001(宏鑫多媒體)	 113.196.120.90-93')
    add_picture('seednet', 'CA008001', 6)

    for i in range(5):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---------------------------------台北SONET IDC機房-------------------------------
    big_title('台北SONET IDC機房')

    # ---------------------------------宏鑫多媒體[1]-------------------------------
    small_title('Traffic Analysis for 宏鑫多媒體[1]	219.84.172.113-117 (6M)')
    add_picture('sonet', '台北_宏鑫多媒體1', 5.32)

    # ---------------------------------宏鑫多媒體[2]-------------------------------
    small_title('Traffic Analysis for 宏鑫多媒體[2]	219.84.172.121-125 (6M)')
    add_picture('sonet', '台北_宏鑫多媒體2', 5.32)

    # ---------------------------------宏鑫多媒體[3]-------------------------------
    small_title('Traffic Analysis for 宏鑫多媒體[3]	219.84.172.145-149 (2M)')
    add_picture('sonet', '台北_宏鑫多媒體3', 5.32)

    # ---------------------------------宏鑫多媒體[4]-------------------------------
    small_title('Traffic Analysis for 宏鑫多媒體[4]  219.84.168.113-125')
    add_picture('sonet', '台北_宏鑫多媒體4', 5.32)

    # ---------------------------------宏鑫多媒體[5]-------------------------------
    small_title('Traffic Analysis for 宏鑫多媒體[5]  219.84.197.17-29')
    add_picture('sonet', '台北_宏鑫多媒體5', 5.32)

    # ---------------------------------中華電信黎明路IDC機房-------------------------------
    big_title('中華電信黎明路IDC機房')

    # ---------------------------------編號：73898925-------------------------------
    small_title('網銀國際 編號：73898925	    60.249.42.224-239 (4M/4M)')
    add_picture('hinet', '73898925', 6)

    # ---------------------------------編號：73806990-------------------------------
    small_title('宏鑫多媒體 編號：73806990      60.249.40.112-127 (4M/4M)')
    add_picture('hinet', '73806990', 6)

    for i in range(8):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---------------------------------網達IDC機房(宏遠)-------------------------------
    big_title('網達IDC機房(宏遠)')

    # ---------------------------------台中_IAP_145.54-------------------------------
    small_title('網銀國際 (9963)台中_IAP_145.54  122.201.145.48-55')
    add_picture('colo', '台中★122.201.145.48-55★9963IAP', 6)

    # ---------------------------------國貿15.246  210.209.15.241-------------------------------
    small_title('網銀國際 (9808)國貿15.246  210.209.15.241')
    add_picture('colo', '國貿★210.209.15.241★9808', 6)

    # ---------------------------------國貿13.198 210.209.13.192-199-------------------------------
    small_title('網銀國際 (500028)國貿13.198 210.209.13.192-199')
    add_picture('colo', '國貿★210.209.13.192-199★500028', 6)

    for i in range(7):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---------------------------------台中-PORT1_3.14-------------------------------
    small_title('網銀國際 (9963)台中-PORT1_3.14  211.73.3.8-15')
    add_picture('colo', '台中★211.73.3.8-15★9963PORT1', 6)

    # ---------------------------------台中-PORT2-7.222-------------------------------
    small_title('網銀國際 (9963)台中-PORT2-7.222  211.73.7.216-223')
    add_picture('colo', '台中★211.73.7.216-223★9963PORT2', 6)

    # ---------------------------------台中-PORT3_7.198-------------------------------
    small_title('網銀國際 (9963)台中-PORT3_7.198  211.73.7.192-198')
    add_picture('colo', '台中★211.73.7.192-198★9963PORT3', 6)

    for i in range(8):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ---------------------------------HK广捷国际控股有限公司機房-------------------------------
    big_title('HK广捷国际控股有限公司機房')

    # ---------------------------------編號：TW43497  203.85.54.80/28-------------------------------
    small_title('網銀國際 編號：TW43497  203.85.54.80/28  TW43497-new2')
    add_picture('hk', 'TW43497-new2', 10)

    # ---------------------------------編號：TW43497  203.85.54.128/29-------------------------------
    small_title('網銀國際 編號：TW43497  203.85.54.128/29  TW43497-new1')
    add_picture('hk', 'TW43497-new1', 10)

    for i in range(5):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ---------------------------------劍俠奇緣3(果核)-------------------------------
    big_title('劍俠奇緣3(果核)')

    # ---------------------------------網銀國際A Total-------------------------------
    small_title('網銀國際A Total (112.121.75.128/28；79三段/27)')
    add_picture('digicentre', '網銀國際A', 6.78)

    # ---------------------------------網銀國際B Total-------------------------------
    small_title('網銀國際B Total (112.121.75.160/27)')
    add_picture('digicentre', '網銀國際B', 6.73)

    # ---------------------------------網銀國際D Total-------------------------------
    small_title('網銀國際D Total (112.121.75.0/27)Po28')
    add_picture('digicentre', '網銀國際D', 7.97)

    for i in range(3):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ---------------------------------中軟IDC機房-------------------------------
    big_title('中軟IDC機房')

    for i in range(len(waninbank_LoginId)):
        add_picture('waninbank', waninbank_LoginId[i], 6)

    try:
        shutil.rmtree('zabbix_screenshot')
        shutil.rmtree('aptg_screenshot')
        shutil.rmtree('colo_screenshot')
        shutil.rmtree('digicentre_screenshot')
        shutil.rmtree('hinet_screenshot')
        shutil.rmtree('hk_screenshot')
        shutil.rmtree('seednet_screenshot')
        shutil.rmtree('waninbank_screenshot')
        shutil.rmtree('sonet_screenshot')
    except:
        print('資料夾刪除失敗!')

    yesterday = datetime.today() + timedelta(-1)
    yesterday_format = yesterday.strftime('%Y-%m.%d')

    doc.save('MRTG統計 ' + yesterday_format + '.docx')

    end_time_now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())
    end_time = datetime.now()
    print('結束時間:', end_time_now)
    print('耗時：', (end_time - start_time).seconds, '秒')



